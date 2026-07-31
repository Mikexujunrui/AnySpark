import html
import io
from datetime import datetime
from typing import cast


def _format_chapter_title(ch: dict, index: int) -> str:
    """Format chapter title, adding [番外] prefix for extra chapters."""
    title = cast(str, ch.get("title", f"第{index}章"))
    if ch.get("is_extra"):
        return f"[番外] {title}"
    return title


def export_txt(book_title: str, chapters: list[dict], include_metadata: bool = False) -> bytes:
    lines = []
    lines.append(f"《{book_title}》")
    lines.append(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    lines.append("")
    lines.append("=" * 40)
    lines.append("")

    for i, ch in enumerate(chapters, 1):
        title = _format_chapter_title(ch, i)
        content = ch.get("content", "")
        lines.append(f"\n{'─' * 20}")
        lines.append(f"  {title}")
        lines.append(f"{'─' * 20}\n")
        if include_metadata:
            lines.append(f"  [创建: {ch.get('createdAt', '未知')}]")
            lines.append(f"  [字数: {len(content)}]")
            lines.append("")
        lines.append(content)
        lines.append("")

    lines.append("")
    lines.append("─── 全书完 ───")
    total = sum(len(ch.get("content", "")) for ch in chapters)
    extra_count = sum(1 for ch in chapters if ch.get("is_extra"))
    stats_parts = [f"共 {len(chapters)} 章，{total} 字"]
    if extra_count:
        stats_parts.append(f"（含番外 {extra_count} 章）")
    lines.append("".join(stats_parts))

    return "\n".join(lines).encode("utf-8")


def export_docx(book_title: str, chapters: list[dict]) -> bytes:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except ImportError:
        raise ImportError("需要安装 python-docx: pip install python-docx")

    doc = Document()

    title_para = doc.add_heading(book_title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"导出时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_page_break()

    extra_count = sum(1 for ch in chapters if ch.get("is_extra"))
    if len(chapters) > 1:
        doc.add_heading("目录", level=1)
        for i, ch in enumerate(chapters, 1):
            title = _format_chapter_title(ch, i)
            doc.add_paragraph(f"{i}. {title}", style="List Number")
        doc.add_page_break()

    for i, ch in enumerate(chapters, 1):
        title = _format_chapter_title(ch, i)
        content = ch.get("content", "")
        doc.add_heading(title, level=1)

        for paragraph in content.split("\n"):
            p = paragraph.strip()
            if p:
                para = doc.add_paragraph()
                para.paragraph_format.first_line_indent = Pt(24)
                run = para.add_run(p)
                run.font.size = Pt(12)

        if i < len(chapters):
            doc.add_page_break()

    doc.add_paragraph("")
    final = doc.add_paragraph("── 全书完 ──")
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    total = sum(len(ch.get("content", "")) for ch in chapters)
    stats_text = f"共 {len(chapters)} 章，{total} 字"
    if extra_count:
        stats_text += f"（含番外 {extra_count} 章）"
    stats = doc.add_paragraph(stats_text)
    stats.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def export_single_chapter_txt(chapter: dict) -> bytes:
    title = chapter.get("title", "章节")
    if chapter.get("is_extra"):
        title = f"[番外] {title}"
    content = chapter.get("content", "")
    text = f"{title}\n\n{content}"
    return text.encode("utf-8")


def export_epub(book_title: str, chapters: list[dict], cover_path: str = "", vertical: bool = False) -> bytes:
    """Export chapters as an EPUB 3.0 ebook.

    Args:
        book_title: The book title.
        chapters: List of chapter dicts with 'title' and 'content'.
        cover_path: Optional path to a cover image file.
        vertical: If True, use vertical writing mode (Chinese vertical text).

    Returns:
        EPUB file content as bytes.
    """
    try:
        from ebooklib import epub
    except ImportError:
        raise ImportError("需要安装 ebooklib: pip install ebooklib")

    book = epub.EpubBook()
    book.set_identifier(f"anyspark-{book_title}")
    book.set_title(book_title)
    book.set_language("zh")
    book.add_author("AnySpark")

    # ── CSS for vertical writing mode ──
    vertical_css = ""
    if vertical:
        vertical_css = """
            html {
                -epub-writing-mode: vertical-rl;
                -webkit-writing-mode: vertical-rl;
                writing-mode: vertical-rl;
            }
        """
    style_css = epub.EpubItem(
        uid="style",
        file_name="style/default.css",
        media_type="text/css",
        content=f"body {{ font-family: serif; line-height: 1.8; }} h2 {{ margin-top: 1em; }} p {{ text-indent: 2em; margin: 0.5em 0; }} {vertical_css}",
    )
    book.add_item(style_css)

    # ── Cover ──
    if cover_path:
        try:
            with open(cover_path, "rb") as f:
                book.set_cover("cover", f.read())
        except Exception:
            pass

    # ── Table of Contents ──
    spine = ["nav"]
    toc = []

    for i, ch in enumerate(chapters, 1):
        title = _format_chapter_title(ch, i)
        content = ch.get("content", "")

        # Convert plain text to HTML paragraphs with XML escaping
        html_lines = [f"<h2>{html.escape(title)}</h2>"]
        for para in content.split("\n"):
            p = para.strip()
            if p:
                html_lines.append(f"<p>{html.escape(p)}</p>")
        html_body = "\n".join(html_lines)

        escaped_title = html.escape(title, quote=True)
        chapter = epub.EpubHtml(
            title=title,
            file_name=f"chap_{i}.xhtml",
            lang="zh",
        )
        chapter.content = (
            f"<!DOCTYPE html>\n"
            f'<html xmlns="http://www.w3.org/1999/xhtml" '
            f'xml:lang="zh">\n'
            f"<head><title>{escaped_title}</title>"
            f'<link rel="stylesheet" type="text/css" href="style/default.css"/>'
            f"</head>\n"
            f"<body>\n{html_body}\n</body>\n</html>"
        )
        book.add_item(chapter)
        spine.append(chapter)
        toc.append(chapter)

    book.toc = toc
    book.spine = spine
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    buffer = io.BytesIO()
    epub.write_epub(buffer, book, {})
    return buffer.getvalue()
